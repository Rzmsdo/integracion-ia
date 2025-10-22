"""
Utilidades para procesamiento y extracción de texto de documentos.
Soporta múltiples formatos: PDF, DOCX, TXT, MD.
"""

import PyPDF2
import docx
import io
from pathlib import Path
from typing import Union
from fastapi import UploadFile


async def extract_text_from_pdf(file_source: Union[str, bytes, UploadFile]) -> str:
    """
    Extrae texto de un archivo PDF.
    
    Args:
        file_source: Ruta del archivo, bytes o UploadFile de FastAPI
    
    Returns:
        Texto extraído del PDF
    """
    try:
        if isinstance(file_source, str):
            # Si es una ruta de archivo
            with open(file_source, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text.strip()
        
        elif isinstance(file_source, bytes):
            # Si son bytes
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_source))
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text.strip()
        
        else:
            # Si es UploadFile de FastAPI
            contents = await file_source.read()
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(contents))
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text.strip()
    
    except Exception as e:
        raise Exception(f"Error extracting text from PDF: {str(e)}")


async def extract_text_from_docx(file_source: Union[str, bytes, UploadFile]) -> str:
    """
    Extrae texto de un archivo DOCX.
    
    Args:
        file_source: Ruta del archivo, bytes o UploadFile de FastAPI
    
    Returns:
        Texto extraído del DOCX
    """
    try:
        if isinstance(file_source, str):
            # Si es una ruta de archivo
            doc = docx.Document(file_source)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text.strip()
        
        elif isinstance(file_source, bytes):
            # Si son bytes
            doc = docx.Document(io.BytesIO(file_source))
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text.strip()
        
        else:
            # Si es UploadFile de FastAPI
            contents = await file_source.read()
            doc = docx.Document(io.BytesIO(contents))
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text.strip()
    
    except Exception as e:
        raise Exception(f"Error extracting text from DOCX: {str(e)}")


async def extract_text_from_txt(file_source: Union[str, bytes, UploadFile]) -> str:
    """
    Lee texto de un archivo de texto plano (TXT, MD, etc.).
    
    Args:
        file_source: Ruta del archivo, bytes o UploadFile de FastAPI
    
    Returns:
        Contenido del archivo de texto
    """
    try:
        if isinstance(file_source, str):
            # Si es una ruta de archivo
            with open(file_source, 'r', encoding='utf-8') as file:
                return file.read().strip()
        
        elif isinstance(file_source, bytes):
            # Si son bytes
            return file_source.decode('utf-8').strip()
        
        else:
            # Si es UploadFile de FastAPI
            contents = await file_source.read()
            return contents.decode('utf-8').strip()
    
    except UnicodeDecodeError:
        # Intentar con otras codificaciones si UTF-8 falla
        try:
            if isinstance(file_source, bytes):
                return file_source.decode('latin-1').strip()
            elif not isinstance(file_source, str):
                contents = await file_source.read()
                return contents.decode('latin-1').strip()
        except Exception as e:
            raise Exception(f"Error decoding text file: {str(e)}")
    
    except Exception as e:
        raise Exception(f"Error reading text file: {str(e)}")


async def extract_text(file_source: Union[str, UploadFile], file_extension: str = None) -> str:
    """
    Extrae texto de un archivo según su extensión.
    Función principal que enruta al extractor apropiado.
    
    Args:
        file_source: Ruta del archivo o UploadFile de FastAPI
        file_extension: Extensión del archivo (opcional, se detecta automáticamente)
    
    Returns:
        Texto extraído del archivo
    """
    # Determinar la extensión
    if file_extension is None:
        if isinstance(file_source, str):
            file_extension = Path(file_source).suffix.lower()
        elif hasattr(file_source, 'filename'):
            file_extension = Path(file_source.filename).suffix.lower()
        else:
            raise ValueError("Cannot determine file extension")
    else:
        file_extension = file_extension.lower()
        if not file_extension.startswith('.'):
            file_extension = f'.{file_extension}'
    
    # Seleccionar el extractor apropiado
    extractors = {
        '.pdf': extract_text_from_pdf,
        '.docx': extract_text_from_docx,
        '.doc': extract_text_from_docx,  # También intentar con .doc
        '.txt': extract_text_from_txt,
        '.md': extract_text_from_txt,
        '.markdown': extract_text_from_txt,
    }
    
    extractor = extractors.get(file_extension)
    if not extractor:
        raise ValueError(
            f"Unsupported file type: {file_extension}. "
            f"Supported types: {', '.join(extractors.keys())}"
        )
    
    return await extractor(file_source)


def get_text_statistics(text: str) -> dict:
    """
    Calcula estadísticas básicas de un texto.
    
    Args:
        text: Texto a analizar
    
    Returns:
        Diccionario con estadísticas
    """
    words = text.split()
    sentences = text.split('.')
    paragraphs = text.split('\n\n')
    
    return {
        "character_count": len(text),
        "word_count": len(words),
        "sentence_count": len(sentences),
        "paragraph_count": len([p for p in paragraphs if p.strip()]),
        "average_word_length": sum(len(word) for word in words) / len(words) if words else 0,
        "average_sentence_length": len(words) / len(sentences) if sentences else 0,
    }


def truncate_text(text: str, max_length: int = 10000) -> str:
    """
    Trunca un texto a una longitud máxima.
    Útil para limitar el input a modelos de IA.
    
    Args:
        text: Texto a truncar
        max_length: Longitud máxima en caracteres
    
    Returns:
        Texto truncado
    """
    if len(text) <= max_length:
        return text
    
    return text[:max_length] + "\n\n[... texto truncado ...]"


def clean_text(text: str) -> str:
    """
    Limpia y normaliza texto extraído.
    Elimina espacios excesivos, líneas vacías, etc.
    
    Args:
        text: Texto a limpiar
    
    Returns:
        Texto limpio
    """
    # Eliminar espacios múltiples
    text = ' '.join(text.split())
    
    # Eliminar líneas vacías múltiples
    lines = text.split('\n')
    cleaned_lines = [line.strip() for line in lines if line.strip()]
    
    return '\n'.join(cleaned_lines)
